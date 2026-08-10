from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
IMG_DIR = ROOT / "ronin_manual_images"
DOCX_PATH = ROOT / "ronin_rig_build_manual.docx"
IMG_DIR.mkdir(parents=True, exist_ok=True)

W, H = 1400, 760
BG = "#F7F3EA"
INK = "#172033"
MUTED = "#667085"
ACCENT = "#F5C84C"
RED = "#A64138"
NAVY = "#20283A"
SKIN = "#A96F4F"
HAIR = "#171726"
PACK = "#5D3527"
CAT = "#8D5139"
WHITE = "#FFFDF7"
TEAL = "#38A892"

FONT_REG = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"

def font(size, bold=False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REG, size)

def rr(d, box, r, fill, outline=None, width=1):
    d.rounded_rectangle(box, radius=r, fill=fill, outline=outline, width=width)

def limb(d, a, b, color, width=34, highlight=False):
    line_color = ACCENT if highlight else color
    d.line([a, b], fill=line_color, width=width)
    rad = width // 2
    for x, y in (a, b):
        d.ellipse((x-rad, y-rad, x+rad, y+rad), fill=line_color)

def draw_piece(d, piece, ox=0, oy=0, highlight=False, ghost=False):
    c = ACCENT if highlight else None
    alpha = "#C9C5BC" if ghost else None
    def pick(base): return c or alpha or base
    if piece == "hair_back":
        d.ellipse((ox+520, oy+108, ox+708, oy+285), fill=pick(HAIR))
        d.polygon([(ox+526,oy+180),(ox+482,oy+228),(ox+550,oy+220)], fill=pick(HAIR))
    elif piece == "leg_back_upper": limb(d,(ox+568,oy+448),(ox+535,oy+535),pick(NAVY),38)
    elif piece == "leg_back_lower": limb(d,(ox+535,oy+535),(ox+575,oy+616),pick(NAVY),34)
    elif piece == "foot_back": rr(d,(ox+548,oy+596,ox+624,oy+632),17,pick(PACK))
    elif piece == "arm_back_upper": limb(d,(ox+575,oy+330),(ox+530,oy+406),pick(NAVY),30)
    elif piece == "arm_back_lower": limb(d,(ox+530,oy+406),(ox+566,oy+470),pick(SKIN),25)
    elif piece == "backpack":
        rr(d,(ox+455,oy+288,ox+548,oy+458),32,pick(PACK))
        d.ellipse((ox+466,oy+252,ox+534,oy+323),fill=pick("#714230"))
    elif piece == "torso":
        rr(d,(ox+535,oy+292,ox+673,oy+468),38,pick(NAVY))
        d.polygon([(ox+540,oy+440),(ox+671,oy+440),(ox+692,oy+495),(ox+526,oy+495)],fill=pick(NAVY))
    elif piece == "head":
        d.ellipse((ox+548,oy+142,ox+710,oy+308), fill=pick(SKIN))
        d.ellipse((ox+674,oy+224,ox+720,oy+260), fill=pick(SKIN))
    elif piece == "leg_front_upper": limb(d,(ox+638,oy+458),(ox+681,oy+534),pick(NAVY),40)
    elif piece == "leg_front_lower": limb(d,(ox+681,oy+534),(ox+652,oy+618),pick(NAVY),35)
    elif piece == "foot_front": rr(d,(ox+635,oy+596,ox+719,oy+634),18,pick(PACK))
    elif piece == "arm_front_upper": limb(d,(ox+654,oy+328),(ox+694,oy+405),pick(NAVY),31)
    elif piece == "arm_front_lower": limb(d,(ox+694,oy+405),(ox+661,oy+472),pick(SKIN),26)
    elif piece == "hair_front":
        d.pieslice((ox+538,oy+125,ox+716,oy+302),180,340,fill=pick(HAIR))
        d.polygon([(ox+575,oy+166),(ox+604,oy+230),(ox+632,oy+163),(ox+662,oy+226),(ox+690,oy+165)],fill=pick(HAIR))
    elif piece == "headband":
        d.line([(ox+548,oy+202),(ox+704,oy+202)], fill=pick(RED), width=17)
        d.polygon([(ox+553,oy+202),(ox+488,oy+180),(ox+526,oy+218)],fill=pick(RED))
    elif piece == "face":
        d.ellipse((ox+645,oy+218,ox+670,oy+254),fill=pick(WHITE))
        d.ellipse((ox+655,oy+228,ox+664,oy+245),fill=pick(INK))
        d.arc((ox+671,oy+259,ox+700,oy+278),10,150,fill=pick(INK),width=3)
    elif piece == "sword":
        d.line([(ox+500,oy+510),(ox+712,oy+405)], fill=pick("#6C7484"), width=20)
        d.line([(ox+684,oy+417),(ox+726,oy+467)], fill=pick(RED), width=17)
        d.line([(ox+692,oy+450),(ox+742,oy+425)], fill=pick(ACCENT if not c else c), width=10)
    elif piece == "cat_body": d.ellipse((ox+190,oy+500,ox+340,oy+582),fill=pick(CAT))
    elif piece == "cat_head":
        d.ellipse((ox+310,oy+468,ox+390,oy+548),fill=pick(CAT))
        d.polygon([(ox+320,oy+483),(ox+328,oy+447),(ox+350,oy+478)],fill=pick(CAT))
        d.polygon([(ox+356,oy+476),(ox+377,oy+447),(ox+382,oy+490)],fill=pick(CAT))
    elif piece == "cat_tail":
        d.arc((ox+145,oy+425,ox+245,oy+558),105,286,fill=pick(CAT),width=22)
    elif piece == "cat_legs":
        limb(d,(ox+225,oy+555),(ox+214,oy+612),pick(CAT),19)
        limb(d,(ox+278,oy+557),(ox+300,oy+612),pick(CAT),19)
        limb(d,(ox+330,oy+545),(ox+360,oy+605),pick(CAT),18)
    elif piece == "cat_face":
        d.ellipse((ox+348,oy+495,ox+365,oy+518),fill=pick(WHITE))
        d.ellipse((ox+354,oy+501,ox+361,oy+512),fill=pick(INK))
        d.ellipse((ox+375,oy+520,ox+383,oy+527),fill=pick(INK))

ORDER = ["hair_back","leg_back_upper","leg_back_lower","foot_back","arm_back_upper","arm_back_lower","backpack","torso","head","leg_front_upper","leg_front_lower","foot_front","arm_front_upper","arm_front_lower","hair_front","headband","face","sword","cat_tail","cat_legs","cat_body","cat_head","cat_face"]

STAGES = [
 ("1", "Set up the tracing frame", [], "Create a 600 × 600 frame. Place the Canva artwork inside, set it to 30% opacity, then lock it."),
 ("2", "Build the rear silhouette", ["hair_back","leg_back_upper","leg_back_lower","foot_back"], "Use simple shapes. Extend each leg segment beneath the joint by about 10–15 px."),
 ("3", "Add the back arm and backpack", ["arm_back_upper","arm_back_lower","backpack"], "Keep the back arm behind the torso. The backpack may overlap the shoulder."),
 ("4", "Add the torso and head", ["torso","head"], "Use one rounded torso shape and one large head ellipse. Keep the head oversized and app-like."),
 ("5", "Build the front leg", ["leg_front_upper","leg_front_lower","foot_front"], "Make thigh, shin and foot separate layers. Hide the top of the thigh beneath the torso."),
 ("6", "Build the front arm", ["arm_front_upper","arm_front_lower"], "Make upper arm and forearm separate. Overlap them at the elbow."),
 ("7", "Add the character details", ["hair_front","headband","face"], "Use only a few bold shapes. Avoid individual hair strands, textures and realistic shading."),
 ("8", "Add the sword", ["sword"], "Keep the complete sword in one group so it can move with the torso or hand."),
 ("9", "Build the cat silhouette", ["cat_tail","cat_legs","cat_body","cat_head"], "Keep the cat in its own group. Separate the tail and legs so they can move."),
 ("10", "Add the cat face", ["cat_face"], "Use only an eye and tiny nose. Readability at phone size matters more than detail."),
]

def state_until(index):
    pieces=[]
    for _,_,added,_ in STAGES[:index+1]: pieces += added
    return pieces

def diagram(index, title, added):
    im=Image.new("RGB",(W,H),BG); d=ImageDraw.Draw(im)
    d.text((54,34),f"STEP {index+1:02}",font=font(25,True),fill=RED)
    d.text((54,72),title,font=font(38,True),fill=INK)
    d.line((54,130,W-54,130),fill="#D8D1C4",width=3)
    if index==0:
        rr(d,(130,185,640,690),28,"#EEE8DD",outline="#B9AE9C",width=4)
        d.text((284,410),"600 × 600",font=font(34,True),fill=MUTED)
        d.text((785,270),"REFERENCE",font=font(24,True),fill=RED)
        d.text((785,312),"Opacity 30%",font=font(31,True),fill=INK)
        d.text((785,365),"Lock layer",font=font(31,True),fill=INK)
        d.text((785,418),"Draw above it",font=font(31,True),fill=INK)
        return im
    previous=state_until(index-1)
    current=state_until(index)
    d.text((70,155),"ADD THESE PIECES",font=font(22,True),fill=MUTED)
    d.text((760,155),"YOUR CANVAS SHOULD NOW LOOK LIKE THIS",font=font(22,True),fill=MUTED)
    rr(d,(64,195,600,700),28,"#EEE8DD")
    rr(d,(726,195,1335,700),28,"#EEE8DD")
    # Exploded miniatures: show new pieces in their real locations, bright yellow.
    for p in previous: draw_piece(d,p,ox=-330,oy=20,ghost=True)
    for p in added: draw_piece(d,p,ox=-330,oy=20,highlight=True)
    # Completed state at a readable scale/location.
    for p in ORDER:
        if p in current: draw_piece(d,p,ox=305,oy=20,highlight=(p in added))
    d.polygon([(650,430),(690,405),(690,420),(715,420),(715,440),(690,440),(690,455)],fill=TEAL)
    d.text((82,653),"New pieces are yellow",font=font(22,True),fill=RED)
    d.text((751,653),"Yellow = check placement, then recolour",font=font(20,True),fill=RED)
    return im

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def add_page_num(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    run=paragraph.add_run("RKA.OS  •  RONIN RIG MANUAL  •  ")
    run.font.name="Arial"; run.font.size=Pt(8); run.font.color.rgb=RGBColor(102,112,133)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); paragraph._p.append(fld)

def add_title(doc,text,size=28,color=INK,after=8):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after)
    r=p.add_run(text); r.bold=True; r.font.name="Arial"; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color[1:])
    return p

def add_body(doc,text,size=11,bold=False,color=INK,after=6):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.18
    r=p.add_run(text); r.bold=bold; r.font.name="Arial"; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color[1:])
    return p

def build_doc():
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Inches(.55); sec.bottom_margin=Inches(.55); sec.left_margin=Inches(.62); sec.right_margin=Inches(.62)
    sec.header_distance=Inches(.25); sec.footer_distance=Inches(.25)
    styles=doc.styles
    normal=styles['Normal']; normal.font.name='Arial'; normal.font.size=Pt(11); normal.font.color.rgb=RGBColor.from_string(INK[1:])
    add_page_num(sec.footer.paragraphs[0])

    # Cover
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(72); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("BUILD MANUAL"); r.bold=True; r.font.name='Arial'; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string(RED[1:])
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("RONIN + CAT\nANIMATION RIG"); r.bold=True; r.font.name='Arial'; r.font.size=Pt(34); r.font.color.rgb=RGBColor.from_string(INK[1:])
    add_body(doc,"A shape-by-shape Figma guide for building a simple, mobile-app mascot that is ready to animate in Rive.",15,False,MUTED,18).alignment=WD_ALIGN_PARAGRAPH.CENTER
    cover=diagram(9,"Finished construction",STAGES[9][2]); cover_path=IMG_DIR/'cover.png'; cover.save(cover_path)
    doc.add_picture(str(cover_path),width=Inches(7.0))
    add_body(doc,"Build rule: simple silhouettes first. Details come only after every moving part exists as its own named layer.",12,True,RED,0).alignment=WD_ALIGN_PARAGRAPH.CENTER

    # Tools page
    doc.add_page_break(); add_title(doc,"Tools — learn these once",28)
    add_body(doc,"You will reuse the same small set of Figma tools throughout the manual.",12,False,MUTED,14)
    tools=[
      ("V — Move / Select","Click a shape to select it. Drag to move. Hold Shift while selecting to choose several shapes."),
      ("O — Ellipse","Press O, then drag. Hold Shift for a perfect circle. Use this for the head, eyes and rounded joints."),
      ("R — Rectangle","Press R and drag. Increase Corner radius in the right panel to make torsos, feet and bags."),
      ("P — Pen","Press P and click around a custom silhouette. Click the first point again to close the shape."),
      ("⌘D — Duplicate","Duplicates the selected shape. Useful for making matching legs, arms and eyes."),
      ("⌘G — Group","Select related pieces and group them. Name groups Ronin, Cat, Sword and Backpack."),
      ("Rename","Double-click a layer name in the left panel. Name it immediately before drawing the next piece."),
      ("Opacity + Lock","Select the reference, set Opacity to 30% in the right panel, then press ⇧⌘L to lock it."),
    ]
    for i,(name,desc) in enumerate(tools,1):
        t=doc.add_table(rows=1,cols=2); t.autofit=False; t.columns[0].width=Inches(1.65); t.columns[1].width=Inches(5.35)
        set_cell_shading(t.cell(0,0),'F5C84C'); set_cell_shading(t.cell(0,1),'F4F0E8')
        a=t.cell(0,0).paragraphs[0]; a.paragraph_format.space_after=Pt(0); rrn=a.add_run(name); rrn.bold=True; rrn.font.name='Arial'; rrn.font.size=Pt(10)
        b=t.cell(0,1).paragraphs[0]; b.paragraph_format.space_after=Pt(0); rrn=b.add_run(desc); rrn.font.name='Arial'; rrn.font.size=Pt(10)
        doc.add_paragraph().paragraph_format.space_after=Pt(1)

    # Construction pages
    for idx,(num,title,added,instruction) in enumerate(STAGES):
        doc.add_page_break(); add_title(doc,f"Step {num} — {title}",25)
        add_body(doc,instruction,12,False,INK,8)
        image=diagram(idx,title,added); path=IMG_DIR/f"step_{idx+1:02}.png"; image.save(path)
        doc.add_picture(str(path),width=Inches(7.1))
        checkpoint = "Checkpoint: " + ("The reference is locked and cannot be selected accidentally." if idx==0 else "Every yellow piece exists as a separate, clearly named layer. Nothing has been flattened.")
        table=doc.add_table(rows=1,cols=1); table.autofit=False; table.columns[0].width=Inches(7.0); set_cell_shading(table.cell(0,0),'E7F5F1')
        p=table.cell(0,0).paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(checkpoint); r.bold=True; r.font.name='Arial'; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string('176B5B')

    # Rigging page
    doc.add_page_break(); add_title(doc,"Final check — prepare for Rive",28)
    add_body(doc,"Do not export until every box below is true.",12,False,MUTED,14)
    checks=[
      "The Ronin and cat are separate top-level groups.",
      "Upper and lower arms are separate layers.",
      "Upper and lower legs are separate layers.",
      "Each joint overlaps by roughly 10–15 px, so rotation cannot reveal gaps.",
      "Back limbs sit behind the torso; front limbs sit in front.",
      "The reference image is hidden before export.",
      "Every layer has a useful name; no important layer is still called Vector or Rectangle.",
      "The character remains readable when viewed at approximately 120 px tall.",
    ]
    for n,item in enumerate(checks,1):
        t=doc.add_table(rows=1,cols=2); t.autofit=False; t.columns[0].width=Inches(.55); t.columns[1].width=Inches(6.45)
        set_cell_shading(t.cell(0,0),'F5C84C'); set_cell_shading(t.cell(0,1),'F4F0E8')
        p=t.cell(0,0).paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(str(n)); r.bold=True; r.font.size=Pt(12)
        p=t.cell(0,1).paragraphs[0]; r=p.add_run(item); r.font.name='Arial'; r.font.size=Pt(11)
        doc.add_paragraph().paragraph_format.space_after=Pt(2)
    add_title(doc,"Recommended layer order",18,RED,6)
    add_body(doc,"Hair Back → Back Leg → Back Arm → Backpack → Torso → Head → Front Leg → Front Arm → Hair Front → Headband → Face → Sword",11,True,INK,8)
    add_body(doc,"Cat: Tail → Back Legs → Body → Front Legs → Head → Face",11,True,INK,8)
    add_body(doc,"When you reach this page, send a screenshot of your Layers panel and canvas. We can correct the structure before you begin animating.",12,True,TEAL,0)

    doc.core_properties.title="Ronin + Cat Animation Rig — Figma Build Manual"
    doc.core_properties.subject="Step-by-step shape construction guide for Rive-ready artwork"
    doc.core_properties.author="RKA.OS"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)

if __name__ == "__main__": build_doc()
