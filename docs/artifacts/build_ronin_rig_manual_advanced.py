from pathlib import Path
from PIL import Image, ImageDraw, ImageFont, ImageFilter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
PROJECT = ROOT.parent.parent
SOURCE = PROJECT / "apps/mobile/assets/ronin/journey/ronin-cat-walkers-v1.png"
IMG_DIR = ROOT / "ronin_manual_advanced_images"
DOCX_PATH = ROOT / "ronin_rig_build_manual_advanced.docx"
IMG_DIR.mkdir(parents=True, exist_ok=True)

W, H = 1500, 820
BG = "#F6F1E7"; PANEL = "#EDE6D9"; INK = "#172033"; MUTED = "#667085"
RED = "#B33A32"; GOLD = "#F2C14E"; TEAL = "#238B78"; BLUE = "#4D78C4"; WHITE = "#FFFDF7"
FONT_REG = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"

def font(size, bold=False): return ImageFont.truetype(FONT_BOLD if bold else FONT_REG, size)
def rr(d, box, radius, fill, outline=None, width=1): d.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)

SRC = Image.open(SOURCE).convert("RGBA")

# Coordinates are in the original 1254 × 1254 source. They intentionally use
# restrained anchor counts; the user should adjust handles, not add dozens of points.
REGIONS = {
 "back_shapes": ((500,220,925,790), [(730,230),(660,245),(624,295),(600,365),(622,480),(585,505),(520,540),(520,700),(610,770),(720,750),(790,645),(825,520),(900,410),(885,285)]),
 "head": ((690,300,925,525), [(727,346),(782,318),(852,330),(900,375),(904,432),(877,485),(820,506),(765,486),(724,440)]),
 "hair_band": ((555,220,920,510), [(590,365),(627,300),(708,244),(790,225),(870,260),(910,325),(888,375),(840,342),(780,370),(710,390),(640,410)]),
 "torso": ((600,475,900,855), [(695,500),(760,485),(832,535),(852,675),(885,800),(810,846),(700,825),(642,720),(645,590)]),
 "arms": ((625,490,920,800), [(674,525),(724,532),(725,662),(695,724),(664,714),(648,642),(650,560)]),
 "legs": ((600,760,950,990), [(690,792),(780,800),(814,842),(792,916),(742,971),(672,965),(625,920),(650,842)]),
 "backpack": ((500,480,730,800), [(577,492),(654,510),(690,590),(677,744),(620,782),(548,758),(520,680),(525,565)]),
 "sword": ((560,585,995,860), [(592,824),(938,598),(976,610),(958,642),(624,852)]),
 "cat_body": ((210,720,505,975), [(260,830),(290,790),(355,785),(420,808),(457,855),(440,916),(370,944),(290,930),(245,890)]),
 "cat_parts": ((210,720,505,985), [(220,825),(228,760),(257,726),(272,752),(260,810),(300,850),(365,840),(410,780),(456,774),(485,816),(472,862),(430,895),(455,956),(410,976),(370,920),(315,916),(305,968),(260,970),(272,900)]),
 "details": ((215,220,985,986), [(215,220),(985,220),(985,986),(215,986)]),
}

STEPS = [
 ("01","Prepare an exact tracing workspace","V, F, opacity, lock","Use a 1000 × 1000 frame. Fit the PNG inside without stretching, set it to 45% opacity and lock it. Create top-level groups named RONIN, CAT and REFERENCE.","details"),
 ("02","Trace the rear silhouette first","Pen tool + Bézier handles","Trace Hair Back, Rear Arm, Rear Leg and the backpack silhouette as separate closed paths. These sit behind the body, so finish their hidden overlap beneath the torso.","back_shapes"),
 ("03","Build the face as a clean skin shape","Pen tool + ellipse","Trace one smooth head/ear silhouette. Use 8–12 anchors and long handles. Add the two eyes as ellipses; do not trace the painted texture or freckles yet.","head"),
 ("04","Construct hair and headband","Pen tool + node editing","Trace Hair Front in 5–7 large locks rather than every strand. Make the headband, knot and two tails separate paths so the tails can sway independently.","hair_band"),
 ("05","Trace the robe in large panels","Pen tool + colour sampler","Build Torso Base, Front Lapel, Rear Lapel, Tunic Skirt and Sash as separate shapes. Preserve the crossed robe silhouette; omit fabric grain.","torso"),
 ("06","Separate both arms at the elbow","Pen tool + rotate check","Create upper arm, forearm and hand layers. Extend each piece 10–15 px underneath its neighbour. Temporarily rotate the forearm 25° to prove the elbow stays covered.","arms"),
 ("07","Separate thighs, shins and boots","Pen tool + duplicate/edit","Trace six moving pieces: front/rear thigh, front/rear shin and two boots. Keep the wide trousers, but use rounded hidden ends at the knee and ankle.","legs"),
 ("08","Rebuild the backpack with controlled depth","Pen, ellipse, boolean union","Use one pack silhouette, one bedroll ellipse/spiral group, two straps and two pockets. Keep the rich brown depth with only a base, shadow and highlight—not the PNG texture.","backpack"),
 ("09","Make the sword independently movable","Pen tool + grouping","Trace sheath, handle, guard and end cap separately, then group them as SWORD. Keep the diagonal angle from the PNG and place it between the back arm and torso.","sword"),
 ("10","Build the cat as its own small rig","Pen tool + ellipse","Trace Body, Head, Muzzle, Tail, four legs and collar separately. Use the PNG silhouette closely; simplify the stripes into five bold shapes.","cat_body"),
 ("11","Refine cat legs, tail and face","Node editing + joint overlap","Give each visible leg a rounded hidden root beneath the body. Keep the tail as one tapered path. Add two eye shapes, muzzle, nose and bell only after the silhouette works.","cat_parts"),
 ("12","Add depth without returning to raster detail","Duplicate, inset paths, clipping","For major pieces, add one shadow shape and one small highlight clipped inside the parent group. Avoid noise, grain and tiny strokes; every detail must remain readable at 120 px tall.","details"),
]

def source_to_box(box, target):
    crop=SRC.crop(box)
    tw,th=target
    crop.thumbnail((tw,th),Image.Resampling.LANCZOS)
    return crop

def map_point(pt, region, dst):
    x0,y0,x1,y1=region; dx,dy,dw,dh=dst
    return (dx+(pt[0]-x0)/(x1-x0)*dw, dy+(pt[1]-y0)/(y1-y0)*dh)

def fit_region(region, panel, pad=26):
    x0,y0,x1,y1=region; px,py,px2,py2=panel; pw,ph=px2-px,py2-py
    rw,rh=x1-x0,y1-y0; scale=min((pw-2*pad)/rw,(ph-2*pad)/rh)
    dw,dh=rw*scale,rh*scale
    return (px+(pw-dw)/2,py+(ph-dh)/2,dw,dh)

def add_anchor_overlay(d, points, region, dst):
    pts=[map_point(p,region,dst) for p in points]
    d.line(pts+[pts[0]],fill=RED,width=5,joint="curve")
    for i,(x,y) in enumerate(pts):
        d.ellipse((x-8,y-8,x+8,y+8),fill=WHITE,outline=RED,width=4)
        if i in (1,len(pts)//2,len(pts)-2):
            hx=x+(28 if i%2 else -28); hy=y-24
            d.line((x,y,hx,hy),fill=BLUE,width=3)
            d.ellipse((hx-5,hy-5,hx+5,hy+5),fill=BLUE)

def diagram(step_index):
    num,title,tool,instruction,key=STEPS[step_index]
    im=Image.new("RGB",(W,H),BG); d=ImageDraw.Draw(im)
    d.text((58,30),f"STEP {num}",font=font(25,True),fill=RED)
    d.text((58,68),title,font=font(38,True),fill=INK)
    d.line((58,126,W-58,126),fill="#D5CDBF",width=3)
    left=(58,168,708,743); right=(790,168,1440,743)
    rr(d,left,30,PANEL); rr(d,right,30,PANEL)
    d.text((82,190),"TRACE / EDIT HERE",font=font(21,True),fill=MUTED)
    d.text((814,190),"THE RESULT MUST STILL MATCH THIS",font=font(21,True),fill=MUTED)
    region,points=REGIONS[key]
    ldst=fit_region(region,left,pad=42); rdst=fit_region(region,right,pad=42)
    crop=SRC.crop(region).resize((int(ldst[2]),int(ldst[3])),Image.Resampling.LANCZOS)
    faded=Image.new("RGBA",crop.size,(255,255,255,0)); faded.alpha_composite(crop); faded.putalpha(faded.getchannel('A').point(lambda a:int(a*.52)))
    im.paste(faded,(int(ldst[0]),int(ldst[1])),faded)
    add_anchor_overlay(d,points,region,ldst)
    crop2=SRC.crop(region).resize((int(rdst[2]),int(rdst[3])),Image.Resampling.LANCZOS)
    im.paste(crop2,(int(rdst[0]),int(rdst[1])),crop2)
    # Arrow and labels
    d.polygon([(778,360),(733,330),(733,348),(705,348),(705,372),(733,372),(733,390)],fill=TEAL)
    d.text((72,752),f"TOOL: {tool}",font=font(22,True),fill=RED)
    d.text((930,752),"Red circles = anchors  •  Blue = Bézier handles",font=font(19,True),fill=BLUE)
    return im

def set_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def add_page_num(p):
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=p.add_run("RKA.OS  •  ADVANCED RONIN VECTOR MANUAL  •  "); r.font.name='Arial'; r.font.size=Pt(8); r.font.color.rgb=RGBColor(102,112,133)
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); p._p.append(fld)

def add_title(doc,text,size=27,color=INK,after=7):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after)
    r=p.add_run(text); r.bold=True; r.font.name='Arial'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color[1:]); return p

def add_body(doc,text,size=11,bold=False,color=INK,after=6):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.18
    r=p.add_run(text); r.bold=bold; r.font.name='Arial'; r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color[1:]); return p

def add_callout(doc,label,text,fill='E7F3EF'):
    t=doc.add_table(rows=1,cols=1); t.autofit=False; t.columns[0].width=Inches(7.0); set_shading(t.cell(0,0),fill)
    p=t.cell(0,0).paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(label+"  "); r.bold=True; r.font.name='Arial'; r.font.size=Pt(10.5); r.font.color.rgb=RGBColor.from_string('176B5B')
    r=p.add_run(text); r.font.name='Arial'; r.font.size=Pt(10.5); r.font.color.rgb=RGBColor.from_string(INK[1:])

def build_doc():
    doc=Document(); sec=doc.sections[0]
    sec.top_margin=Inches(.52); sec.bottom_margin=Inches(.5); sec.left_margin=Inches(.62); sec.right_margin=Inches(.62); sec.header_distance=Inches(.25); sec.footer_distance=Inches(.24)
    doc.styles['Normal'].font.name='Arial'; doc.styles['Normal'].font.size=Pt(11)
    add_page_num(sec.footer.paragraphs[0])
    # Cover
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(42); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("ADVANCED BUILD MANUAL"); r.bold=True; r.font.name='Arial'; r.font.size=Pt(14); r.font.color.rgb=RGBColor.from_string(RED[1:])
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("RONIN + CAT\nVECTOR RIG"); r.bold=True; r.font.name='Arial'; r.font.size=Pt(35); r.font.color.rgb=RGBColor.from_string(INK[1:])
    add_body(doc,"Trace the real artwork closely while rebuilding every moving part as a clean, editable Figma path.",14,False,MUTED,14).alignment=WD_ALIGN_PARAGRAPH.CENTER
    cover=Image.new('RGB',(1400,720),BG); source=SRC.copy(); source.thumbnail((850,650),Image.Resampling.LANCZOS); cover.paste(source,((1400-source.width)//2,20),source)
    cp=IMG_DIR/'advanced_cover.png'; cover.save(cp); doc.add_picture(str(cp),width=Inches(7.0))
    add_callout(doc,"OUTCOME","A recognisable version of the supplied PNG—not a generic substitute—structured for Rive animation.")

    # Tool primer page 1
    doc.add_page_break(); add_title(doc,"Advanced tools — paths and curves",28)
    add_body(doc,"Learn these once. Every construction page then tells you only which tool to use.",12,False,MUTED,12)
    tools=[
      ("P — Pen: straight points","Click once for a corner. Use corners only where the silhouette genuinely changes direction."),
      ("P — Pen: curved points","Click and drag while placing a point to pull Bézier handles. Longer handles create smoother curves."),
      ("Enter — Edit object","Select a vector and press Enter. Drag anchors or handles. Double-click an anchor to switch between corner and smooth."),
      ("Delete an anchor","While editing a path, select an unnecessary point and press Delete. Fewer anchors usually produce a cleaner result."),
      ("Close a path","With Pen active, click the first anchor again. Closed paths can receive fills and import reliably into Rive."),
      ("I — Eyedropper","Select a shape, press I and sample the reference. Use sampled colours as a starting point, then simplify them."),
    ]
    for name,desc in tools:
        t=doc.add_table(rows=1,cols=2); t.autofit=False; t.columns[0].width=Inches(1.72); t.columns[1].width=Inches(5.28)
        set_shading(t.cell(0,0),'F2C14E'); set_shading(t.cell(0,1),'F2EDE4')
        p=t.cell(0,0).paragraphs[0]; r=p.add_run(name); r.bold=True; r.font.size=Pt(10)
        p=t.cell(0,1).paragraphs[0]; r=p.add_run(desc); r.font.size=Pt(10)
        doc.add_paragraph().paragraph_format.space_after=Pt(2)
    add_callout(doc,"ANCHOR RULE","Start with one point at each major change in direction. Adjust handles before adding another point.", 'E8EEF8')

    # Tool primer page 2
    doc.add_page_break(); add_title(doc,"Advanced tools — construction and depth",28)
    tools2=[
      ("⌘D — Duplicate then edit","Duplicate a related limb or detail, then reshape its anchors. Do not mirror blindly; the PNG uses perspective."),
      ("Boolean Union","Select overlapping shapes and choose Union from the toolbar when they should become one silhouette."),
      ("Use as Mask","Place a detail over its parent, select both and choose Use as Mask. Prefer keeping masks simple for the Rive handoff."),
      ("Rotation test","Move the transform origin mentally to the joint, rotate a limb 20–30°, check for gaps, then Undo."),
      ("Outline view: ⇧O","Toggle outline view to spot accidental open paths, duplicate points and overly complicated silhouettes."),
      ("Scale check","Zoom out until the combined character is about 120 px tall. Remove any detail that becomes visual noise."),
    ]
    for name,desc in tools2:
        t=doc.add_table(rows=1,cols=2); t.autofit=False; t.columns[0].width=Inches(1.72); t.columns[1].width=Inches(5.28)
        set_shading(t.cell(0,0),'F2C14E'); set_shading(t.cell(0,1),'F2EDE4')
        p=t.cell(0,0).paragraphs[0]; r=p.add_run(name); r.bold=True; r.font.size=Pt(10)
        p=t.cell(0,1).paragraphs[0]; r=p.add_run(desc); r.font.size=Pt(10)
        doc.add_paragraph().paragraph_format.space_after=Pt(2)
    add_title(doc,"Fidelity target",18,RED,6)
    add_body(doc,"Preserve: silhouette, proportions, large hair masses, robe crossover, red headband/sash, backpack roll, sword angle, cat silhouette and expressive eyes.",11,True,INK,6)
    add_body(doc,"Simplify: fabric grain, tiny scratches, freckles, hair strand texture, repeated stitching and photographic shading.",11,True,MUTED,6)
    add_callout(doc,"DEPTH RULE","Use at most three tones per major material: base, shadow and a restrained highlight.")

    # Step pages
    for idx,(num,title,tool,instruction,key) in enumerate(STEPS):
        doc.add_page_break(); add_title(doc,f"Step {num} — {title}",25)
        add_body(doc,instruction,11.5,False,INK,7)
        img=diagram(idx); path=IMG_DIR/f"advanced_step_{idx+1:02}.png"; img.save(path); doc.add_picture(str(path),width=Inches(7.08))
        checkpoints={
          0:"The reference is locked; RONIN and CAT groups are empty and ready above it.",
          1:"Rear pieces extend behind future torso shapes, with no visible gaps at the shoulder or hip.",
          2:"The face silhouette matches the PNG at 100% zoom and remains smooth in outline view.",
          3:"Headband tails, hair back and hair front are independent layers.",
          4:"The robe crossover remains clearly recognisable when all texture is hidden.",
          5:"Both forearms can rotate 25° without opening a hole at either elbow.",
          6:"Thighs, shins and boots rotate independently without losing the baggy-trouser silhouette.",
          7:"The backpack reads as deep brown leather with no more than three tonal levels.",
          8:"SWORD is one named group and does not share vector points with the hand or torso.",
          9:"The cat reads correctly as a silhouette before stripes or face details are visible.",
          10:"Tail and legs move independently; face details remain legible at phone size.",
          11:"At 120 px tall the design still resembles the supplied PNG, with no noisy micro-detail.",
        }
        add_callout(doc,"CHECK BEFORE CONTINUING",checkpoints[idx])

    # Final structure
    doc.add_page_break(); add_title(doc,"Final layer structure and Rive gate",28)
    add_body(doc,"Match this organisation before exporting. The exact number of decorative shapes may differ; moving parts must remain separate.",11.5,False,MUTED,10)
    columns=[
      ("RONIN — behind torso",["Hair Back","Rear Upper Arm","Rear Forearm","Rear Hand","Rear Thigh","Rear Shin","Rear Boot","Backpack","Sword / Sheath"]),
      ("RONIN — body/front",["Torso Base","Robe Panels","Head / Skin","Hair Front","Headband + Tails","Front Upper Arm","Front Forearm","Front Hand","Front Thigh","Front Shin","Front Boot","Face Details","Sash"]),
      ("CAT",["Tail","Rear Legs","Body","Front Legs","Head","Muzzle","Eyes + Nose","Collar + Bell","Stripes"]),
    ]
    for title,items in columns:
        add_title(doc,title,16,RED,3)
        add_body(doc,"  •  ".join(items),10.5,True,INK,8)
    gates=["All visible paths are closed","No moving limb shares a path with the torso","Every elbow, knee, shoulder and hip has hidden overlap","Reference PNG is hidden","No unnamed Vector layers remain","Character passes the 120 px scale check","A duplicate Figma frame is kept as a backup before Rive import"]
    for i,g in enumerate(gates,1):
        t=doc.add_table(rows=1,cols=2); t.autofit=False; t.columns[0].width=Inches(.5); t.columns[1].width=Inches(6.5)
        set_shading(t.cell(0,0),'F2C14E'); set_shading(t.cell(0,1),'F2EDE4')
        p=t.cell(0,0).paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(str(i)); r.bold=True
        p=t.cell(0,1).paragraphs[0]; r=p.add_run(g); r.font.size=Pt(10.5)
        doc.add_paragraph().paragraph_format.space_after=Pt(1)
    add_callout(doc,"NEXT","Send a screenshot of the canvas plus the complete Layers panel before importing into Rive. Structural mistakes are much easier to fix in Figma.")

    doc.core_properties.title="Advanced Ronin + Cat Vector Rig Build Manual"
    doc.core_properties.subject="Faithful Pen-tool reconstruction of the supplied PNG for Rive"
    doc.core_properties.author="RKA.OS"
    doc.save(DOCX_PATH); print(DOCX_PATH)

if __name__ == '__main__': build_doc()
